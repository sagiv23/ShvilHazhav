import os
from docx import Document
from docx.shared import Pt, RGBColor
from pygments.lexers import JavaLexer
from pygments.token import Token

"""
Utility script to export all Java source files from the project into a single Microsoft Word document.
This is often used for academic submissions or code reviews.

Dependencies:
- python-docx: install via 'pip install python-docx'
- pygments: install via 'pip install pygments'
"""

# מיפוי מורחב של טוקנים לצבעים המדויקים של Android Studio (IntelliJ Light Theme)
TOKEN_STYLES = {
    # מילות מפתח (כחול כהה)
    Token.Keyword: {'color': RGBColor(0x00, 0x33, 0xB3), 'bold': True},
    Token.Keyword.Type: {'color': RGBColor(0x00, 0x33, 0xB3), 'bold': True},
    Token.Keyword.Declaration: {'color': RGBColor(0x00, 0x33, 0xB3), 'bold': True},

    # מחרוזות (ירוק)
    Token.String: {'color': RGBColor(0x06, 0x7D, 0x17)},
    Token.Literal.String: {'color': RGBColor(0x06, 0x7D, 0x17)},

    # הערות (אפור נטוי)
    Token.Comment: {'color': RGBColor(0x8C, 0x8C, 0x8C), 'italic': True},
    Token.Comment.Single: {'color': RGBColor(0x8C, 0x8C, 0x8C), 'italic': True},
    Token.Comment.Multiline: {'color': RGBColor(0x8C, 0x8C, 0x8C), 'italic': True},

    # מספרים (כחול בהיר)
    Token.Number: {'color': RGBColor(0x17, 0x50, 0xEB)},

    # שמות פונקציות/מתודות
    Token.Name.Function: {'color': RGBColor(0x00, 0x62, 0x7A)},

    # מחלקות (שחור מודגש)
    Token.Name.Class: {'color': RGBColor(0x00, 0x00, 0x00), 'bold': True},

    # אנוטציות כמו @Override או @Nullable (זהוב/חום)
    Token.Name.Decorator: {'color': RGBColor(0x9E, 0x88, 0x0D)},

    # משתני מחלקה / Fields (סגול)
    Token.Name.Variable: {'color': RGBColor(0x87, 0x10, 0x94)},
    Token.Name.Attribute: {'color': RGBColor(0x87, 0x10, 0x94)},
}

def get_style_for_token(token_type):
    """מחזיר את העיצוב המתאים לסוג הטוקן, כולל בדיקה של סוגי אב."""
    if token_type in TOKEN_STYLES:
        return TOKEN_STYLES[token_type]

    # בדיקה בהיררכיה של הטוקנים (למשל Token.Keyword.Declaration -> Token.Keyword)
    parent = token_type.parent
    while parent:
        if parent in TOKEN_STYLES:
            return TOKEN_STYLES[parent]
        parent = parent.parent
    return {}

def export_java_to_word(project_path, output_docx):
    """
    Crawls the project directory and appends the content of every .java file to a Word document.
    """
    doc = Document()
    doc.add_heading('Shvil Hazhav - Java Source Code', 0)

    # ניתוב לתיקיית המקור של ה-Java כדי שהנתיבים יהיו קצרים ויחסיים לחבילה (Package)
    java_base_path = os.path.join(project_path, "app", "src", "main", "java", "com", "example", "sagivproject")

    # בדיקה שהתיקייה קיימת, אם לא - נשתמש בנתיב הפרויקט הכללי
    search_path = java_base_path if os.path.exists(java_base_path) else project_path

    # Walk through the project directory structure
    for root, dirs, files in os.walk(search_path):
        # Skip hidden directories like .git or .idea
        if any(hidden in root for hidden in [".git", ".idea", "build", "gradle"]):
            continue

        for file in files:
            if file.endswith(".java"):
                file_path = os.path.join(root, file)
                # חישוב נתיב יחסי מהחבילה הראשית
                relative_path = os.path.relpath(file_path, search_path)
                # החלפת לוכסנים לפורמט ווינדוס והוספת הסימן מהדוגמה
                display_path = "▶ " + relative_path.replace('/', '\\')

                print(f"Processing: {display_path}")

                # הוספת כותרת ועיצובה (Arial, 12, #09890E, ללא Bold)
                heading = doc.add_heading('', level=1)
                run = heading.add_run(display_path)
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = False  # ביטול ההדגשה (Bold)
                run.font.color.rgb = RGBColor(0x09, 0x89, 0x0E)

                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        code_content = f.read()

                    # הוספת הקוד עם פונט מונוספייס והדגשת תחביר (Syntax Highlighting)
                    paragraph = doc.add_paragraph()
                    lexer = JavaLexer()

                    for ttype, value in lexer.get_tokens(code_content):
                        run = paragraph.add_run(value)
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)

                        style = get_style_for_token(ttype)
                        if 'color' in style:
                            run.font.color.rgb = style['color']
                        if style.get('bold'):
                            run.font.bold = True
                        if style.get('italic'):
                            run.font.italic = True

                except Exception as e:
                    doc.add_paragraph(f"Error reading file {display_path}: {e}")

                doc.add_page_break()

    doc.save(output_docx)
    print(f"\nOperation complete! Source code saved to: {output_docx}")

if __name__ == "__main__":
    # Define the project root relative to this script's location
    PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
    OUTPUT_FILE = "ShvilHazhav_SourceCode_Export.docx"

    export_java_to_word(PROJECT_ROOT, OUTPUT_FILE)